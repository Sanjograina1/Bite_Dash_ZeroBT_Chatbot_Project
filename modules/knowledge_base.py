"""
knowledge_base.py — Document extraction, chunking, embedding, and BM25 indexing.
"""

import io
import re
from pathlib import Path

import numpy as np
from openai import OpenAI
from rank_bm25 import BM25Okapi

from modules import store

KB_DIR = Path(__file__).parent.parent / "knowledge_base"

# ── Text Extraction ──────────────────────────────────────────────────

def extract_text_with_pages(file_path: str, file_type: str) -> list[dict]:
    """Return list of {page: int, text: str} dicts."""
    pages: list[dict] = []

    if file_type == "pdf":
        pages = _extract_pdf(file_path)
    elif file_type in {"docx", "doc"}:
        pages = _extract_docx(file_path)
    else:
        pages = _extract_txt(file_path)

    return pages


def extract_any_file(file_bytes: bytes, filename: str, api_key: str = "") -> str:
    """Extract readable text content from ANY file format (PDF, DOCX, TXT, CSV, Excel, Images, Audio, Code, etc.)."""
    ext = Path(filename).suffix.lower()

    # 1. Plain text / Code / Structured Data files
    text_exts = {
        ".txt", ".md", ".csv", ".json", ".xml", ".html", ".py", ".js", ".ts", ".jsx", ".tsx",
        ".css", ".log", ".yaml", ".yml", ".ini", ".cfg", ".sql", ".sh", ".rtf", ".env"
    }
    if ext in text_exts or not ext:
        for enc in ["utf-8", "latin-1", "cp1252", "ascii"]:
            try:
                return file_bytes.decode(enc)
            except Exception:
                continue

    # 2. PDF files
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = [p.extract_text() for p in pdf.pages if p.extract_text()]
                if pages_text:
                    return "\n\n".join(pages_text)
        except Exception:
            pass
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = [p.extract_text() for p in reader.pages if p.extract_text()]
            if pages_text:
                return "\n\n".join(pages_text)
        except Exception:
            pass

    # 3. DOCX files
    if ext in {".docx", ".doc"}:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        except Exception:
            pass

    # 4. Excel spreadsheet files (.xlsx, .xls)
    if ext in {".xlsx", ".xls"}:
        try:
            import pandas as pd
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
            sheets_text = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheets_text.append(f"--- Sheet: {sheet_name} ---\n" + df.to_string(index=False))
            return "\n\n".join(sheets_text)
        except Exception:
            pass

    # 5. Audio files (Transcribe using Whisper API)
    if ext in {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac"}:
        if api_key:
            try:
                client = OpenAI(api_key=api_key)
                transcript = client.audio.transcriptions.create(
                    model="whisper-1", file=(filename, file_bytes)
                )
                return f"[Audio Transcription of {filename}]:\n" + transcript.text
            except Exception as e:
                return f"[Audio File: {filename} (size: {len(file_bytes)} bytes)]"

    # 6. Image files (Vision OCR analysis via GPT-4o)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}:
        if api_key:
            try:
                import base64
                client = OpenAI(api_key=api_key)
                b64_img = base64.b64encode(file_bytes).decode("utf-8")
                mime_type = "image/png" if ext == ".png" else "image/jpeg"
                resp = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": f"Extract and summarize all visible text, numbers, table data, and visual details from this image ({filename}):"},
                            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64_img}"}}
                        ]
                    }],
                    max_tokens=1500
                )
                return f"[Image Analysis & Vision Text Extraction of {filename}]:\n" + resp.choices[0].message.content.strip()
            except Exception as e:
                return f"[Image File: {filename} (size: {len(file_bytes)} bytes)]"

    # 7. General fallback for any other arbitrary file type
    for enc in ["utf-8", "latin-1"]:
        try:
            return file_bytes.decode(enc, errors="ignore")[:4000]
        except Exception:
            pass

    return f"[File Attachment: {filename} ({len(file_bytes)} bytes)]"


def _extract_pdf(path: str) -> list[dict]:
    try:
        import pdfplumber
        result = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                t = page.extract_text()
                if t and t.strip():
                    result.append({"page": i, "text": t})
        if result:
            return result
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        return [
            {"page": i + 1, "text": p.extract_text()}
            for i, p in enumerate(reader.pages)
            if p.extract_text() and p.extract_text().strip()
        ]
    except Exception:
        return []


def _extract_docx(path: str) -> list[dict]:
    try:
        from docx import Document
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [{"page": 1, "text": full_text}] if full_text else []
    except Exception:
        return []


def _extract_txt(path: str) -> list[dict]:
    try:
        text = Path(path).read_text(encoding="utf-8", errors="ignore")
        return [{"page": 1, "text": text}] if text.strip() else []
    except Exception:
        return []


# ── Chunking with Metadata ──────────────────────────────────────────

def chunk_with_metadata(
    pages: list[dict],
    source_name: str,
    chunk_size: int = 400,
    overlap: int = 80,
) -> list[dict]:
    """Sliding-window chunking that preserves source file and page number."""
    chunks: list[dict] = []
    for page_info in pages:
        page_num = page_info["page"]
        text = page_info["text"]
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            snippet = text[start:end].strip()
            if snippet:
                chunks.append({
                    "text": snippet,
                    "source_file": source_name,
                    "page_number": page_num,
                })
            start += chunk_size - overlap
    return chunks


# ── Embedding ────────────────────────────────────────────────────────

def embed_texts(texts: list[str], api_key: str) -> np.ndarray:
    """Embed a list of texts using OpenAI text-embedding-3-small."""
    client = OpenAI(api_key=api_key)
    vectors: list[list[float]] = []
    batch_size = 100
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        resp = client.embeddings.create(model="text-embedding-3-small", input=batch)
        for item in resp.data:
            vectors.append(item.embedding)
    return np.array(vectors, dtype=np.float32)


# ── Full Ingest Pipeline ────────────────────────────────────────────

def add_document(uploaded_file, api_key: str, chunk_size: int = 400,
                 overlap: int = 80) -> int:
    """Process an uploaded file end-to-end and return the doc_id."""
    KB_DIR.mkdir(exist_ok=True)

    name = uploaded_file.name
    suffix = name.rsplit(".", 1)[-1].lower()
    save_path = KB_DIR / name

    # Save to disk
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Check and replace existing active document with same name
    existing_docs = store.list_documents()
    for d in existing_docs:
        if d["name"] == name and d["status"] == "active":
            store.soft_delete_document(d["id"])

    # Register in DB
    doc_id = store.add_document_record(name, suffix)

    # Extract + chunk
    pages = extract_text_with_pages(str(save_path), suffix)
    if not pages:
        store.soft_delete_document(doc_id)
        raise ValueError(f"Could not extract text from {name}")

    chunks = chunk_with_metadata(pages, name, chunk_size, overlap)

    # Embed
    texts = [c["text"] for c in chunks]
    vectors = embed_texts(texts, api_key)

    # Prepare for DB
    for i, ch in enumerate(chunks):
        ch["embedding"] = vectors[i].tobytes()

    store.save_chunks(doc_id, chunks)
    store.update_document_chunks(doc_id, len(chunks))

    return doc_id


# ── Load Everything ──────────────────────────────────────────────────

def load_index(api_key: str | None = None):
    """Load all active chunks, embeddings, and build BM25 index.
    Returns (chunks, vectors, bm25) or ([], None, None) if empty."""
    raw = store.load_active_chunks()
    if not raw:
        return [], None, None

    chunks = []
    vecs = []
    for r in raw:
        chunks.append({
            "id": r["id"],
            "text": r["text"],
            "source_file": r["source_file"],
            "page_number": r["page_number"],
        })
        vecs.append(np.frombuffer(r["embedding"], dtype=np.float32))

    vectors = np.stack(vecs)

    # BM25 index
    tokenized = [_tokenize(c["text"]) for c in chunks]
    bm25 = BM25Okapi(tokenized) if tokenized else None

    return chunks, vectors, bm25


def _tokenize(text: str) -> list[str]:
    """Simple word tokenization for BM25."""
    return re.findall(r"\w+", text.lower())
